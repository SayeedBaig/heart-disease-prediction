from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "docs" / "major_project_defense_guide.md"
OUT = ROOT / "docs" / "major_project_defense_guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 40)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_diagram(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    text = "\n".join(lines).strip()
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    add_run(p, text, color=INK, size=8.5, font="Consolas")
    doc.add_paragraph()


def add_callout(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.2)
    text = " ".join(line.strip("> ").strip() for line in lines).strip()
    add_run(p, text, italic=True, color=INK, size=10.5)
    doc.add_paragraph()


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "Major Project Defense Guide", color=MUTED, size=9)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    quote = styles["Intense Quote"]
    quote.font.name = "Calibri"
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = INK


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=48, after=8, line=1.1)
    add_run(title, "Major Project Defense Guide", bold=True, color=BLUE, size=24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=0, after=18, line=1.1)
    add_run(
        subtitle,
        "A Multi-Model Multi-Stage Framework for Heart Disease Prediction Using Intelligent Feature Fusion and Risk Modelling",
        color=INK,
        size=13,
    )

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Current Status", "Phase-1 Completed; Phase-2 In Progress"),
        ("Primary Use", "Placement interview, viva, and project defense preparation"),
        ("Language Style", "Simple English with terminology meanings"),
        ("Source References", "Instruction PDF, Review-1 PPT, Review-2 PPT, speaker notes, and project codebase"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_fill(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            else:
                set_cell_fill(cell, "FFFFFF")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        note,
        "This document separates team contribution, individual contribution, current implementation, and future Phase-2 work.",
        italic=True,
        color=MUTED,
        size=10.5,
    )
    doc.add_page_break()


def add_table_of_contents_note(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=12, line=1.2)
    add_run(
        p,
        "Navigation note: Use the Word navigation pane to jump through headings. The document is structured with clear section headings for interview preparation.",
        italic=True,
        color=MUTED,
        size=10,
    )


def render_markdown(doc, text):
    in_code = False
    code_lines = []
    quote_lines = []

    def flush_quote():
        nonlocal quote_lines
        if quote_lines:
            add_callout(doc, quote_lines)
            quote_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            add_diagram(doc, code_lines)
            code_lines = []

    for raw in text.splitlines():
        line = raw.rstrip()
        if line.strip() == "---":
            flush_quote()
            flush_code()
            continue
        if line.startswith("```"):
            flush_quote()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith(">"):
            quote_lines.append(line)
            continue
        flush_quote()

        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=0, after=4, line=1.15)
            add_run(p, line[2:].strip(), color=INK, size=10.8)
            continue
        if line[:3].strip(".").isdigit() and ". " in line[:5]:
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, before=0, after=4, line=1.15)
            add_run(p, line.split(". ", 1)[1].strip(), color=INK, size=10.8)
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=6, line=1.25)
        if line.startswith("Q:"):
            add_run(p, "Q:", bold=True, color=DARK_BLUE)
            add_run(p, line[2:])
        elif line.startswith("A:"):
            add_run(p, "A:", bold=True, color=DARK_BLUE)
            add_run(p, line[2:])
        elif line.startswith("Follow-up:"):
            add_run(p, "Follow-up:", bold=True, color=DARK_BLUE)
            add_run(p, line[len("Follow-up:"):])
        else:
            add_run(p, line, color=INK, size=11)

    flush_quote()
    flush_code()


def main():
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_table_of_contents_note(doc)
    render_markdown(doc, text)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
